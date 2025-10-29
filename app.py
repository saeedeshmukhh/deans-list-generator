import streamlit as st
import pandas as pd
from docx import Document
import os, tempfile, zipfile, io, shutil, subprocess

# --------------------------------------------------
# PAGE CONFIG
# --------------------------------------------------
st.set_page_config(page_title="Dean's List Certificates Generator", page_icon="🏅", layout="centered")
st.title("🏅 Dean's List Certificates Generator")

st.write("""
Upload your **Excel file** (Student Data) and **Word template** to generate
personalized Dean’s List letters or certificates for each student.
The app will automatically generate PDFs if your system supports it,
otherwise it will generate Word documents (.docx) instead.
""")

# --------------------------------------------------
# FILE UPLOADS
# --------------------------------------------------
uploaded_excel = st.file_uploader("📄 Upload Student Data Sheet", type=["xlsx"])
uploaded_template = st.file_uploader("📝 Upload Word Template", type=["docx"])

# --------------------------------------------------
# DETECT AVAILABLE PDF ENGINE
# --------------------------------------------------
def detect_pdf_engine():
    # Prefer LibreOffice if available
    if shutil.which("soffice"):
        return "libreoffice"
    # Otherwise, try pypandoc + pandoc
    try:
        import pypandoc
        pypandoc.get_pandoc_version()
        return "pandoc"
    except Exception:
        return None

PDF_ENGINE = detect_pdf_engine()
if PDF_ENGINE:
    st.success(f"PDF engine detected: {PDF_ENGINE}")
else:
    st.warning("⚠️ No PDF engine detected (LibreOffice or Pandoc). "
               "Files will be generated as DOCX instead of PDF.")

# --------------------------------------------------
# MAIN PROCESS
# --------------------------------------------------
if uploaded_excel and uploaded_template:
    if st.button("🚀 Generate Files"):
        with st.spinner("Generating personalized files... ⏳"):
            with tempfile.TemporaryDirectory() as td:
                excel_path = os.path.join(td, "students.xlsx")
                template_path = os.path.join(td, "template.docx")
                open(excel_path, "wb").write(uploaded_excel.read())
                open(template_path, "wb").write(uploaded_template.read())

                df = pd.read_excel(excel_path, engine="openpyxl")

                output_dir = os.path.join(td, "output")
                os.makedirs(output_dir, exist_ok=True)

                def generate_doc(student):
                    doc = Document(template_path)
                    for p in doc.paragraphs:
                        for key, value in student.items():
                            placeholder = f"[{key.upper()}]"
                            if placeholder in p.text:
                                for run in p.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(value))
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for key, value in student.items():
                                    placeholder = f"[{key.upper()}]"
                                    if placeholder in cell.text:
                                        cell.text = cell.text.replace(placeholder, str(value))

                    filename = f"{student['NAME']}_DeansList.docx"
                    filepath = os.path.join(output_dir, filename)
                    doc.save(filepath)
                    return filepath

                def convert_to_pdf(filepath):
                    """Try to convert DOCX to PDF using available engine"""
                    pdf_path = os.path.splitext(filepath)[0] + ".pdf"
                    if PDF_ENGINE == "libreoffice":
                        tmpdir = tempfile.mkdtemp()
                        subprocess.run(
                            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, filepath],
                            check=True
                        )
                        temp_pdf = os.path.join(tmpdir, os.path.basename(pdf_path))
                        shutil.move(temp_pdf, pdf_path)
                        shutil.rmtree(tmpdir, ignore_errors=True)
                    elif PDF_ENGINE == "pandoc":
                        import pypandoc
                        pypandoc.convert_file(filepath, "pdf", outputfile=pdf_path)
                    else:
                        return None
                    return pdf_path

                generated_files = []
                for _, row in df.iterrows():
                    docx_file = generate_doc(row.to_dict())
                    if PDF_ENGINE:
                        try:
                            pdf_file = convert_to_pdf(docx_file)
                            if pdf_file:
                                generated_files.append(pdf_file)
                                continue
                        except Exception as e:
                            st.warning(f"PDF conversion failed for {row['NAME']}: {e}")
                    generated_files.append(docx_file)

                # Zip results
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    for file in generated_files:
                        zf.write(file, arcname=os.path.basename(file))
                zip_buffer.seek(0)

                st.success("✅ All files generated successfully!")
                st.download_button(
                    label="⬇️ Download All Files (ZIP)",
                    data=zip_buffer,
                    file_name="DeansList_Files.zip",
                    mime="application/zip"
                )
else:
    st.info("Please upload both the Excel file and Word template to proceed.")